import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = "/Users/babitakironvedantam/Desktop/CAPSTONE FINAL/GIT hub"
IMG_DIR = os.path.join(BASE_DIR, "outputs/doc_diagrams")
DOCX_OUT = os.path.join(BASE_DIR, "PharmaTrace_AI_Architecture_and_Technical_Specification.docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    run = h.runs[0]
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(11, 44, 94) # Deep Navy
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(14, 116, 144) # Cyan/Teal
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 65, 85) # Slate
    return h

def add_callout(doc, title, text, border_hex="0284c7", bg_hex="f0f9ff"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(f"{title}\n")
    r_t.bold = True
    r_t.font.size = Pt(9.5)
    r_t.font.color.rgb = RGBColor(15, 23, 42)
    
    r_body = p.add_run(text)
    r_body.font.size = Pt(9)
    r_body.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def style_table(tbl, col_widths, headers, data):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    # Headers
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "0f172a") # Dark Slate
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row in enumerate(data):
        row_cells = tbl.add_row().cells
        bg = "f8fafc" if r_idx % 2 == 1 else "ffffff"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].width = col_widths[c_idx]
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(30, 41, 59)
            
    doc_add_tbl_borders(tbl)

def doc_add_tbl_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="0" w:color="cbd5e1"/><w:bottom w:val="single" w:sz="8" w:space="0" w:color="94a3b8"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

print("Helper functions defined successfully.")
